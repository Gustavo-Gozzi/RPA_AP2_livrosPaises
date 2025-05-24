from database import paises
from tabelaPaises import Paises
from modelpais import gerenciarpaises
from database import livros
from tabelaLivros import Livros
from modellivros import gerenciarlivros

from docx.shared import Inches
import requests
from io import BytesIO
import datetime
from docx import Document

gerenciarpaises()
gerenciarlivros()

doc = Document()

integrantes = {
    "Gustavo": {
        "nome": "Gustavo Ranéa Gozzi",
        "ra": "2400779",
        "apelido": "GustavoGozzi"
    },
    "Joao": {
        "nome": "João Vitor de Matos Gouveia",
        "ra": "2400784",
        "apelido": "JoaoVitor"
    }
}

doc.add_heading('Relatório', 0)
doc.add_paragraph(f"""Integrantes do Grupo: 
{integrantes["Gustavo"]["nome"]} - {integrantes["Gustavo"]["ra"]}
{integrantes["Joao"]["nome"]} - {integrantes["Joao"]["ra"]}""")
doc.add_paragraph(f'Data de geração: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")}\n')

doc.add_heading('Países', level=1)

with paises.app_context():
    paises = Paises.query.all()
    for idx, pais in enumerate(paises, 1):
        doc.add_heading(f"{idx}. {pais.nome_comum}", level=2)

        doc.add_paragraph(f"Nome Oficial: {pais.nome_oficial}")
        doc.add_paragraph(f"Capital: {pais.capital}")
        doc.add_paragraph(f"Região: {pais.regiao}")
        doc.add_paragraph(f"Sub-região: {pais.subregiao}")
        doc.add_paragraph(f"População: {pais.populacao:,}")
        doc.add_paragraph(f"Moeda Símbolo: {pais.moedaSimbolo}")
        doc.add_paragraph(f"Moeda Nome: {pais.moedaNome}")
        doc.add_paragraph(f"Idioma: {pais.idioma}")
        doc.add_paragraph(f"Fuso Horário: {pais.fusoHorario}")

        doc.add_paragraph("Bandeira:")

        try:
            response = requests.get(pais.bandeira)
            if response.status_code == 200:
                image = BytesIO(response.content)
                doc.add_picture(image, width=Inches(2))  # Tamanho ajustável
            else:
                doc.add_paragraph("Imagem não disponível.")
        except Exception as e:
            doc.add_paragraph(f"Erro ao baixar imagem: {e}")

        doc.add_paragraph("")

doc.add_heading('2. Dados dos Livros', level=1)

with livros.app_context():
    livros = Livros.query.all()
    for idx, livro in enumerate(livros, 1):
        doc.add_heading(f"{idx}. {livro.titulo}", level=2)

        doc.add_paragraph(f"Preço: {livro.preco}")
        doc.add_paragraph(f"Avaliação: {livro.avaliacao}")
        doc.add_paragraph(f"Disponibilidade: {livro.disponibilidade}")
        doc.add_paragraph("")

data_hora = datetime.datetime.now().strftime("%d-%m-%Y_%H-%M")
doc.save(f'relatorio_{integrantes["Gustavo"]["apelido"]}_{integrantes["Joao"]["apelido"]}_{data_hora}.docx')

print("Relatório gerado com sucesso!")
